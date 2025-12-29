from pathlib import Path
from typing import Optional

from .md_to_html import markdown_to_html
from .html_to_md import html_to_markdown


class ExportError(RuntimeError):
    pass


def export_pdf(html_content: str, output_path: Path, css_text: Optional[str] = None) -> Path:
    try:
        from fpdf import FPDF, HTMLMixin
    except Exception as e:
        raise ExportError("fpdf2 is required") from e

    output_path.parent.mkdir(parents=True, exist_ok=True)

    class PDF(FPDF, HTMLMixin):
        pass

    pdf = PDF()
    pdf.add_page()

    import os
    import re
    font_path = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'msyh.ttc')
    font_path_bd = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'msyhbd.ttc')

    if os.path.exists(font_path):
        pdf.add_font('msyh', '', font_path)
        pdf.add_font('msyh', 'B', font_path_bd if os.path.exists(font_path_bd) else font_path)
        pdf.add_font('msyh', 'I', font_path)
        pdf.add_font('msyh', 'BI', font_path_bd if os.path.exists(font_path_bd) else font_path)
        pdf.set_font('msyh', size=12)
        font_family = 'msyh'
    else:
        pdf.set_font('Helvetica', size=12)
        font_family = 'Helvetica'

    # 移除 code/pre 标签，避免切换到 courier 字体导致中文报错
    html_clean = re.sub(r'<code[^>]*>(.*?)</code>', r'\1', html_content, flags=re.DOTALL)
    html_clean = re.sub(r'<pre[^>]*>(.*?)</pre>', r'<p>\1</p>', html_clean, flags=re.DOTALL)

    pdf.write_html(html_clean, font_family=font_family)
    pdf.output(str(output_path))
    return output_path


def export_docx_from_html(html_content: str, output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise ExportError("python-docx is required for DOCX export. Install with pip install python-docx") from e

    import re
    from bs4 import BeautifulSoup

    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')

    def add_runs(para, element):
        """递归处理内联元素，保留粗体/斜体"""
        if element.name is None:  # 文本节点
            para.add_run(str(element))
        elif element.name in ('strong', 'b'):
            run = para.add_run(element.get_text())
            run.bold = True
        elif element.name in ('em', 'i'):
            run = para.add_run(element.get_text())
            run.italic = True
        elif element.name == 'code':
            run = para.add_run(element.get_text())
            run.font.name = 'Consolas'
        elif element.name == 'a':
            para.add_run(element.get_text())
        else:
            for child in element.children:
                add_runs(para, child)

    def process_element(el):
        if el.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(el.name[1])
            doc.add_heading(el.get_text(), level=level)
        elif el.name == 'p':
            para = doc.add_paragraph()
            for child in el.children:
                add_runs(para, child)
        elif el.name in ('ul', 'ol'):
            for i, li in enumerate(el.find_all('li', recursive=False)):
                prefix = f"{i+1}. " if el.name == 'ol' else "• "
                para = doc.add_paragraph(prefix + li.get_text())
        elif el.name == 'table':
            rows = el.find_all('tr')
            if rows:
                cols = max(len(r.find_all(['td', 'th'])) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.rows[i].cells[j].text = cell.get_text().strip()
        elif el.name == 'pre':
            code = el.get_text()
            para = doc.add_paragraph()
            run = para.add_run(code)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif el.name == 'blockquote':
            para = doc.add_paragraph(el.get_text())
            para.paragraph_format.left_indent = Inches(0.5)
        elif el.name == 'hr':
            doc.add_paragraph('─' * 50)

    for el in soup.find_all(['h1','h2','h3','h4','h5','h6','p','ul','ol','table','pre','blockquote','hr']):
        process_element(el)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_content(
    content: str,
    content_type: str,
    export_format: str,
    target_path: Path,
    css_text: Optional[str] = None,
) -> Path:
    """
    content_type: 'html' or 'md'
    export_format: 'pdf' or 'docx'
    """
    if content_type == "md":
        html_content = markdown_to_html(content)
    else:
        html_content = content

    if export_format == "pdf":
        return export_pdf(html_content, target_path.with_suffix(".pdf"), css_text=css_text)
    elif export_format == "docx":
        return export_docx_from_html(html_content, target_path.with_suffix(".docx"))
    else:
        raise ExportError(f"Unsupported export format: {export_format}")
